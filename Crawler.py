from selenium.webdriver.common.by import By
from ast import keyword
from xml.dom.expatbuilder import TEXT_NODE
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from urllib.parse import quote_plus
from urllib.request import urlopen
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches
import os
import datetime
import sys, getopt

# When there are no folder 
def folder_prob(directory):
    try:
        if not os.path.exists(directory):
            os.makedirs(directory)
    except OSError:
        print('Error: Creating directory. ' + directory)

# Saving images
def save_imgs(imgs, num, keyword, save_path):
    for index, image in enumerate(imgs[:num]):  # images[:크롤링하고 싶은 사진 개수]
        src = image.get_attribute('src')
        t = urlopen(src).read()
        time = datetime.datetime.today().strftime ('%Y%m%d')
        file = open(os.path.join(save_path, time + "_" + keyword + str(index + 1) + ".jpg"), "wb")
        file.write(t)
        print("img save " + save_path + time + "_" + keyword + str(index + 1) + ".jpg")


def main(argv):
    FILE_NAME = argv[0] # command line arguments의 첫번째는 파일명
    SITE = "google"
    KEYWORD = ""
    EXCLUDE = "ㅈ"
    TYPE = "excel"
    NUM = 10
    SAVE_PATH = os.getcwd() + "/"

    try:
        opts, etc_args = getopt.getopt(argv[1:], \
        "hs:k:e:t:n:sp:", ["help","site=","keyword=", "exclude=", "type=", "number=", "savepath="])

    except getopt.GetoptError: # 옵션지정이 올바르지 않은 경우
        print(FILE_NAME, '-s <site> -k <keyword>')
        sys.exit(2)

    for opt, arg in opts: # 옵션이 파싱된 경우
        if opt in ("-h", "--help"): # HELP 요청인 경우 사용법 출력
            print(FILE_NAME, '-s <site> -k <keyword>')
            sys.exit()

        elif opt in ("-s", "--site"):
            SITE = arg

        elif opt in ("-k", "--keyword"):
            KEYWORD = arg

        elif opt in ("-e", "--exclude"):
            EXCLUDE = arg

        elif opt in ("-t", "--type"):
            TYPE = arg

        elif opt in ("-n", "--number"):
            NUM = arg

        elif opt in ("-sp", "--savepath"):
            SAVE_PATH = arg

    if len(KEYWORD) < 1: # 필수항목 값이 비어있다면
        print(FILE_NAME, "-k option is mandatory") # 필수임을 출력
        sys.exit(2)

    print("SITE:",  SITE)
    print("KEYWORD:", KEYWORD)
    print("EXCLUDE:",  EXCLUDE)
    print("TYPE:",  TYPE)
    print("NUM:",  NUM)
    print("SAVE_PATH:",  SAVE_PATH)

    crawling(SITE, KEYWORD, EXCLUDE, TYPE, NUM, SAVE_PATH)

def crawling(site, keyword, exclude, type, num, save_path):
    import time
    num = int(num)

    if site == "google":
        headline = "WlydOe"
        pgNum = 0
        pgNum = str(pgNum)
        url = "http://www.google.com/search?q=" + keyword + "&biw=746&bih=722&tbm=nws&start=" + pgNum

        if type == "img":
            url ="https://images.google.com/search?q=" + keyword + "&tbm=isch&"
            img = "rg_i"

    elif site == "naver":
        headline ="news_tit"
        pgNum = 1
        pgNum = str(pgNum)
        url = "https://search.naver.com/search.naver?where=news&sm=tab_jum&query=" + keyword + "&start=" + pgNum

        if type == "img":
            url ="https://search.naver.com/search.naver?where=image&section=image&query=" + keyword
            img = "_image"

    else:
        headline = "tit_main"
        pgNum = 0
        pgNum = str(pgNum)
        url = "https://search.daum.net/search?w=news&DA=PGD&enc=utf&cluster_page=1&q=" + keyword + "&p=" + pgNum

        if type == "img":
            url ="https://search.daum.net/search?w=img&nil_search=btn&DA=NTB&enc=utf8&q=" + keyword
            img = "thumb_img"

    # Opening Browser
    webdriver_options = webdriver.ChromeOptions()
    webdriver_options .add_argument('--headless')
    webdriver_options .add_argument('--no-sandbox')
    chromedriver = os.getcwd() + "/chromedriver"
    driver = webdriver.Chrome(chromedriver, options=webdriver_options )

    driver.implicitly_wait(3)
    driver.get(url)

    # Input is Doc
    if type == "doc":
        document = Document()
        document.add_heading(keyword, 0)
        
        if num > 20:
            num += 10

        if site == "google":
            count = 0
            while num > 0:
                headlines = driver.find_elements(By.CLASS_NAME, headline)
                titles = driver.find_elements(By.CSS_SELECTOR, 'div[role="heading"][aria-level="3"]')

                for i, elem in enumerate(headlines[:num]):
                    text = titles[i].get_attribute('innerText')
                    href = elem.get_attribute('href')
                    exclude = str(exclude)
                    if exclude in text:
                        continue
                    document.add_paragraph(text)
                    document.add_paragraph(href) 
                    count += 1

                num -= count
                pgNum = int(pgNum)
                pgNum += 10
                pgNum = str(pgNum)
                url = "http://www.google.com/search?q=" + keyword + "&biw=746&bih=722&tbm=nws&start=" + pgNum
                driver.get(url)
                time.sleep(3)

        elif site == "naver":
            count = 0
            while num > 0:
                headlines = driver.find_elements(By.CLASS_NAME, headline)

                for i in headlines[:num]:
                    text = i.text
                    text = str(text)
                    href = i.get_attribute('href')
                    exclude = str(exclude)
                    if exclude in text:
                        continue
                    document.add_paragraph(text)
                    document.add_paragraph(href)
                    count += 1

                num -= count
                pgNum = int(pgNum)
                pgNum += 10
                pgNum = str(pgNum)
                url = "https://search.naver.com/search.naver?where=news&sm=tab_jum&query=" + keyword + "&start=" + pgNum
                driver.get(url)
                time.sleep(3)

        else:
            count = 0
            pgNum = 1 
            while num > 0:
                headlines = driver.find_elements(By.CLASS_NAME, headline)

                for i in headlines[:num]:
                    text = i.text
                    text = str(text)
                    href = i.get_attribute('href')
                    exclude = str(exclude)
                    if exclude in text:
                        continue
                    document.add_paragraph(text)
                    document.add_paragraph(href)
                    count += 1

                num -= count
                pgNum = int(pgNum)
                pgNum += 1
                pgNum = str(pgNum)
                url = "https://search.daum.net/search?w=img&nil_search=btn&DA=NTB&enc=utf8&q=" + keyword + "&p=" + pgNum
                driver.get(url)
                time.sleep(3)

        time = datetime.datetime.today().strftime ('%Y%m%d')
        document.save(save_path + time + "_" + keyword + ".docx")

    # Input is excel
    elif type == "excel":
        wb = Workbook()
        ws = wb.active
        ws.append(["Title", "URL"])
        count = 0
        
        if num > 20:
            num += 10

        if site == "google":
            while num > 0:
                headlines = driver.find_elements(By.CLASS_NAME, headline)
                titles = driver.find_elements(By.CSS_SELECTOR, 'div[role="heading"][aria-level="3"]')

                for i, elem in enumerate(headlines[:num]):
                    text = titles[i].get_attribute('innerText')
                    href = elem.get_attribute('href')
                    exclude = str(exclude)
                    if exclude in text:
                        continue
                    count += 1
                    ws.append([text, href])

                num -= count
                pgNum = int(pgNum)
                pgNum += 10
                pgNum = str(pgNum)
                url = "http://www.google.com/search?q=" + keyword + "&biw=746&bih=722&tbm=nws&start=" + pgNum
                driver.get(url)
                time.sleep(3)

        elif site == "naver":
            count = 0
            while num > 0:
                headlines = driver.find_elements(By.CLASS_NAME, headline)

                for i in headlines[:num]:
                    text = i.text
                    href = i.get_attribute('href')
                    exclude = str(exclude)
                    if exclude in text:
                        continue
                    count += 1
                    ws.append([text, href])

                num -= count
                pgNum = int(pgNum)
                pgNum += 10
                pgNum = str(pgNum)
                url = "https://search.naver.com/search.naver?where=news&sm=tab_jum&query=" + keyword + "&start=" + pgNum
                driver.get(url)
                time.sleep(3) 

        else:
            count = 0
            pgNum = 1
            while num > 0:
                headlines = driver.find_elements(By.CLASS_NAME, headline)

                for i in headlines[:num]:
                    text = i.text
                    href = i.get_attribute('href')
                    exclude = str(exclude)
                    if exclude in text:
                        continue
                    count += 1
                    ws.append([text, href])

                num -= count
                pgNum = int(pgNum)
                pgNum += 1
                pgNum = str(pgNum)
                url = "https://search.daum.net/search?w=img&nil_search=btn&DA=NTB&enc=utf8&q=" + keyword + "&p=" + pgNum
                driver.get(url)
                time.sleep(3) 

        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter # Get the column name

            for cell in col:
                try: # Necessary to avoid error on empty cells
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column].width = adjusted_width
        time = datetime.datetime.today().strftime ('%Y%m%d')
        time = str(time)       
        wb.save(save_path + time + "_" + keyword + ".xlsx")

    # Input is image
    else:
        imgs = driver.find_elements(By.CLASS_NAME, img)
        folder_prob(save_path)
        save_imgs(imgs, num, keyword, save_path)

    driver.close()

if __name__ == "__main__":
    main(sys.argv)
